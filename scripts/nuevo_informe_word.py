"""Scaffold de informes Word: crea el módulo Python + la plantilla .docx base.

Uso:
    python scripts/nuevo_informe_word.py <nombre> [--label "Mi Informe"] [--force]

Crea (si no existen):
    backend/rgenerator/reports/word/informes/<nombre>.py
    backend/rgenerator/reports/word/templates/<nombre>.docx

La plantilla base es un Word real editable que INDICA los términos a
reemplazar con códigos {{valor}} — el usuario la abre en Word, ajusta
estilos/textos manteniendo los códigos, y el módulo Python define qué
valor va en cada código vía `construir_contexto`.
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
WORD_PKG = ROOT / "backend" / "rgenerator" / "reports" / "word"
INFORMES_DIR = WORD_PKG / "informes"
TEMPLATES_DIR = WORD_PKG / "templates"

STUB_MODULO = '''"""Informe Word: {label}.

Editar `construir_contexto` para definir qué valor va en cada código
{{{{valor}}}} de la plantilla `templates/{nombre}.docx`.

Ver ejemplos en `resumen_indicador.py` y helpers en `..engine`:
    - tabla_desde_df(df)  → filas para loops {{%tr for %}}
    - Grafico(fn=..., df=..., params=...) → imagen matplotlib del CHART_REGISTRY
    - Imagen(path=...)    → imagen existente en disco
"""
from __future__ import annotations

from datetime import date

import pandas as pd

from ..engine import Grafico, Imagen, tabla_desde_df  # noqa: F401

LABEL = "{label}"
DESCRIPCION = ""
PARAMS_ESPERADOS = ["titulo"]


def construir_contexto(dataframes: dict[str, pd.DataFrame], params: dict) -> dict:
    """Devuelve el dict Jinja para la plantilla.

    Args:
        dataframes: {{rol: DataFrame}} del indicador (ej "estudiantes").
        params: parámetros enviados desde el frontend.
    """
    return {{
        "titulo": params.get("titulo", "{label}"),
        "fecha": date.today().strftime("%d/%m/%Y"),
        # TODO: agregar aquí los valores para cada código de la plantilla
    }}
'''


def crear_plantilla_base(path: Path, label: str) -> None:
    """Genera un .docx base con los códigos {{valor}} de ejemplo."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    doc.add_heading("{{ titulo }}", level=1)
    doc.add_paragraph("{{ subtitulo }}")
    doc.add_paragraph("Fecha: {{ fecha }}")
    doc.add_paragraph("Registros: {{ n_registros }} — Columnas: {{ columnas }}")

    guia = doc.add_paragraph()
    run = guia.add_run(
        "GUÍA (borrar en la versión final): esta plantilla se edita en Word. "
        "Todo texto entre llaves dobles es un código que el módulo Python "
        f"informes/{path.stem}.py reemplaza al generar el informe. "
        "Las tablas dinámicas usan etiquetas tr-for / tr-endfor en filas "
        "propias (ver la tabla de ejemplo). Las imágenes se insertan donde "
        "esté su código. OJO: no escribir sintaxis de etiquetas en texto "
        "normal — Jinja intentaría interpretarla."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x33, 0x00)
    run.italic = True

    doc.add_heading("Resumen por categoría", level=2)
    tabla = doc.add_table(rows=4, cols=3)
    tabla.style = "Table Grid"
    hdr = tabla.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Categoría", "N", "Promedio"
    tabla.rows[1].cells[0].text = "{%tr for fila in resumen %}"
    fila = tabla.rows[2].cells
    fila[0].text = "{{ fila.categoria }}"
    fila[1].text = "{{ fila.n }}"
    fila[2].text = "{{ fila.promedio }}"
    tabla.rows[3].cells[0].text = "{%tr endfor %}"

    doc.add_heading("Gráfico", level=2)
    doc.add_paragraph("{{ grafico_logro }}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("nombre", help="identificador del informe (snake_case, será el filename)")
    parser.add_argument("--label", default=None, help="nombre legible para el frontend")
    parser.add_argument("--force", action="store_true", help="sobreescribir si existe")
    args = parser.parse_args()

    nombre = args.nombre.strip().lower().replace("-", "_").replace(" ", "_")
    if not nombre.isidentifier():
        print(f"ERROR: '{nombre}' no es un identificador Python válido")
        return 1
    label = args.label or nombre.replace("_", " ").title()

    modulo_path = INFORMES_DIR / f"{nombre}.py"
    plantilla_path = TEMPLATES_DIR / f"{nombre}.docx"

    if modulo_path.exists() and not args.force:
        print(f"= módulo ya existe (se conserva): {modulo_path.relative_to(ROOT)}")
    else:
        modulo_path.write_text(
            STUB_MODULO.format(nombre=nombre, label=label), encoding="utf-8"
        )
        print(f"+ módulo creado:    {modulo_path.relative_to(ROOT)}")

    if plantilla_path.exists() and not args.force:
        print(f"= plantilla ya existe (se conserva): {plantilla_path.relative_to(ROOT)}")
    else:
        crear_plantilla_base(plantilla_path, label)
        print(f"+ plantilla creada: {plantilla_path.relative_to(ROOT)}")

    print(
        f"\nSiguientes pasos:\n"
        f"  1) Editar la plantilla en Word manteniendo los códigos {{{{valor}}}}\n"
        f"  2) Definir los valores en {modulo_path.relative_to(ROOT)}\n"
        f"  3) Probar: POST /api/reports/word/{nombre} con {{indicator_id, filtros, params}}"
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
